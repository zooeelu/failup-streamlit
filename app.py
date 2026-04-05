import os
import json
import re
from datetime import datetime
from io import BytesIO
from pathlib import Path
import xml.etree.ElementTree as ET

import pdfplumber
import requests
from docx import Document
from flask import Flask, send_file, request, jsonify
from anthropic import Anthropic
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)

client = Anthropic(
    api_key=os.getenv("ANTHROPIC_API_KEY")
)

DATA_DIR = Path("data")
DATA_DIR.mkdir(exist_ok=True)

SUBMISSIONS_FILE = DATA_DIR / "submissions.json"
TRIALS_FILE = DATA_DIR / "ctgov_trials.json"
PUBMED_FILE = DATA_DIR / "pubmed_articles.json"

CTGOV_API_URL = "https://clinicaltrials.gov/api/v2/studies"
NCBI_EUTILS_BASE = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"

NCBI_TOOL = os.getenv("NCBI_TOOL", "failup")
NCBI_EMAIL = os.getenv("NCBI_EMAIL", "")
NCBI_API_KEY = os.getenv("NCBI_API_KEY", "")


# ---------------------------
# File helpers
# ---------------------------

def load_json_file(path):
    if not path.exists():
        return []
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
            return data if isinstance(data, list) else []
    except Exception:
        return []


def save_json_file(path, data):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


def load_submissions():
    return load_json_file(SUBMISSIONS_FILE)


def save_submissions(submissions):
    save_json_file(SUBMISSIONS_FILE, submissions)


def load_trials():
    return load_json_file(TRIALS_FILE)


def save_trials(trials):
    save_json_file(TRIALS_FILE, trials)


def load_pubmed_articles():
    return load_json_file(PUBMED_FILE)


def save_pubmed_articles(records):
    save_json_file(PUBMED_FILE, records)


# ---------------------------
# General helpers
# ---------------------------

def now_iso():
    return datetime.utcnow().isoformat() + "Z"


def generate_doi(submission_count):
    return f"10.5281/fu.{1000000 + submission_count}"


def normalize_text(value):
    if not value:
        return ""
    return str(value).strip().lower()


def safe_int(value, default):
    try:
        return int(value)
    except Exception:
        return default


def truncate_label(text, n=24):
    text = text or ""
    return text if len(text) <= n else text[:n] + "…"


def safe_title(value, fallback="Untitled"):
    value = (value or "").strip()
    return value if value else fallback


def sanitize_graph_value(value, fallback="unspecified"):
    value = (value or "").strip()
    return value if value else fallback


# ---------------------------
# Upload extraction helpers
# ---------------------------

ALLOWED_UPLOAD_EXTENSIONS = {".pdf", ".doc", ".docx"}
MAX_EXTRACTED_TEXT_CHARS = 25000


def allowed_upload_file(filename):
    return Path(filename or "").suffix.lower() in ALLOWED_UPLOAD_EXTENSIONS


def extract_text_from_pdf_bytes(file_bytes):
    chunks = []
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                chunks.append(page_text.strip())
    return "\n\n".join(chunks).strip()


def extract_text_from_docx_bytes(file_bytes):
    doc = Document(BytesIO(file_bytes))
    parts = []

    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join((cell.text or "").strip() for cell in row.cells if (cell.text or "").strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts).strip()


def extract_text_from_uploaded_file(storage):
    filename = storage.filename or "uploaded_file"
    suffix = Path(filename).suffix.lower()
    file_bytes = storage.read()
    storage.stream.seek(0)

    if suffix == ".pdf":
        return extract_text_from_pdf_bytes(file_bytes)
    if suffix in {".doc", ".docx"}:
        return extract_text_from_docx_bytes(file_bytes)

    raise ValueError(f"Unsupported file type: {suffix}")


def build_uploaded_document_context(files):
    documents = []

    for storage in files:
        filename = storage.filename or "uploaded_file"
        if not allowed_upload_file(filename):
            raise ValueError(f"Unsupported file type for {filename}. Please upload PDF or DOCX files.")

        extracted = extract_text_from_uploaded_file(storage)
        if extracted:
            documents.append(
                f"FILE: {filename}\n"
                f"BEGIN_CONTENT\n{extracted[:MAX_EXTRACTED_TEXT_CHARS]}\nEND_CONTENT"
            )

    return "\n\n".join(documents)


def heuristic_extract_fields(text):
    clean_text = text or ""
    compact_text = re.sub(r"\s+", " ", clean_text).strip()
    lines = [line.strip() for line in clean_text.splitlines() if line.strip()]

    nct_match = re.search(r"\bNCT\d{8}\b", clean_text, flags=re.IGNORECASE)
    sample_match = re.search(r"(?:sample size|enrolled|enrollment|randomized|participants|subjects)\D{0,15}(\d{2,5})", clean_text, flags=re.IGNORECASE)
    planned_match = re.search(r"(?:planned|target)\s+(?:sample size|enrollment|n)\D{0,15}(\d{2,5})", clean_text, flags=re.IGNORECASE)
    pval_match = re.search(r"\bp\s*[=<>]\s*0?\.?\d+[\d]*", clean_text, flags=re.IGNORECASE)
    hr_match = re.search(r"\b(HR|OR|RR|IRR|MD|SMD|β)\s*[:=]?\s*[^\n.;]{1,80}", clean_text, flags=re.IGNORECASE)

    title = ""
    for line in lines[:12]:
        if len(line) >= 20 and not re.match(r"^(abstract|background|methods|results|introduction)\b", line, flags=re.IGNORECASE):
            title = line
            break

    return {
        "title": title,
        "nct": nct_match.group(0).upper() if nct_match else "",
        "study_type": "",
        "phase": "",
        "n": sample_match.group(1) if sample_match else "",
        "nplan": planned_match.group(1) if planned_match else "",
        "design": "",
        "therapeutic_area": "",
        "intervention": "",
        "population": "",
        "dropout": "",
        "female": "",
        "age": "",
        "result_type": "",
        "outcome": "",
        "pval": pval_match.group(0) if pval_match else "",
        "effect_estimate_type": hr_match.group(1).upper() if hr_match else "",
        "effect_estimate_other": "",
        "effect_estimate": hr_match.group(0) if hr_match else "",
        "limitations": compact_text[:400] if compact_text else ""
    }


def extract_study_fields_with_claude(document_text, file_names):
    fallback = heuristic_extract_fields(document_text)

    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1200,
        system=(
            "You extract structured study metadata from uploaded research documents for a submission form. "
            "Return only valid raw JSON with the exact requested keys. "
            "Do not invent details. If a field is not clearly supported by the documents, return an empty string. "
            "Use concise values that can be inserted directly into HTML form fields. "
            "When possible, map values to the form's existing choices exactly.\n"
            "Allowed study_type values: Randomized clinical trial; Non-randomized controlled trial; Observational (cohort); Observational (case-control); Cross-sectional; Preclinical (in vivo); Preclinical (in vitro); Systematic review / meta-analysis; Basket / platform trial.\n"
            "Allowed phase values: Phase III; Phase II/III; Phase II; Phase I/II; Phase I; Phase IV; Not applicable.\n"
            "Allowed design values: Double-blind RCT — 1:1; Double-blind RCT — other ratio; Single-blind RCT; Open-label RCT; Non-randomized; Not applicable.\n"
            "Allowed therapeutic_area values: Oncology; Neurology; Cardiology; Immunology; Hepatology; Psychiatry; Infectious disease; Endocrinology; Respiratory; Rare disease; Other.\n"
            "Allowed result_type values: null; inconclusive; opposite.\n"
            "Allowed effect_estimate_type values: HR; OR; RR; RD; MD; SMD; β; IRR; NNT; other."
        ),
        messages=[
            {
                "role": "user",
                "content": f"""
Return ONLY valid JSON with exactly these keys:
title
nct
study_type
phase
n
nplan
design
therapeutic_area
intervention
population
dropout
female
age
result_type
outcome
pval
effect_estimate_type
effect_estimate_other
effect_estimate
limitations

Instructions:
- Use empty strings for anything not clearly stated.
- Keep intervention, population, outcome, and limitations concise but informative.
- For n and nplan, return only the value as text that fits the form field.
- For female and age, return concise display text such as "46%" or "58 years".
- For pval, keep the original statistical notation if reported.
- For effect_estimate, include the estimate and confidence interval if available.
- If the estimate type is not one of the allowed values, set effect_estimate_type to "other" and fill effect_estimate_other.

Uploaded files: {", ".join(file_names)}

Fallback guesses from regex (use only if supported by document text):
{json.dumps(fallback, ensure_ascii=False)}

Document text:
{document_text}
"""
            }
        ]
    )

    data = safe_parse_json(response.content[0].text)

    expected_keys = [
        "title", "nct", "study_type", "phase", "n", "nplan", "design",
        "therapeutic_area", "intervention", "population", "dropout", "female",
        "age", "result_type", "outcome", "pval", "effect_estimate_type",
        "effect_estimate_other", "effect_estimate", "limitations"
    ]

    normalized = {}
    for key in expected_keys:
        value = data.get(key, fallback.get(key, ""))
        normalized[key] = value.strip() if isinstance(value, str) else str(value or "").strip()

    if normalized["result_type"] not in {"null", "inconclusive", "opposite"}:
        normalized["result_type"] = fallback.get("result_type", "")

    allowed_effect_types = {"", "HR", "OR", "RR", "RD", "MD", "SMD", "β", "IRR", "NNT", "other"}
    if normalized["effect_estimate_type"] not in allowed_effect_types:
        normalized["effect_estimate_other"] = normalized["effect_estimate_type"]
        normalized["effect_estimate_type"] = "other"

    return normalized


# ---------------------------
# Source labels
# ---------------------------

SOURCE_LABELS = {
    "manual": "Manual submission",
    "ctgov_imported": "Imported CT.gov",
    "ctgov_enriched": "AI-tagged CT.gov",
    "ctgov_curated": "Curated failure record (CT.gov)",
    "pubmed_imported": "Imported PubMed",
    "pubmed_enriched": "AI-tagged PubMed",
    "pubmed_curated": "Curated failure record (PubMed)"
}


def classify_trial_source(trial):
    curation = trial.get("curation", {}) or {}
    if curation.get("review_status") == "curated_failure" and curation.get("graph_worthy") is True:
        return "ctgov_curated"
    if trial.get("summary"):
        return "ctgov_enriched"
    return "ctgov_imported"


def classify_pubmed_source(article):
    curation = article.get("curation", {}) or {}
    if curation.get("review_status") == "curated_failure" and curation.get("graph_worthy") is True:
        return "pubmed_curated"
    if article.get("summary"):
        return "pubmed_enriched"
    return "pubmed_imported"


# ---------------------------
# Heuristics
# ---------------------------

def infer_trial_result_type(trial):
    """
    Prototype heuristic only.
    """
    status = normalize_text(trial.get("status", ""))

    if "terminated" in status or "withdrawn" in status:
        return "opposite"
    if "suspended" in status:
        return "inconclusive"
    if "completed" in status:
        return "null"
    return "inconclusive"


NEGATIVE_TITLE_HINTS = [
    "negative trial",
    "null result",
    "no benefit",
    "no improvement",
    "failed",
    "failure",
    "did not improve",
    "did not meet",
    "not associated",
    "lack of efficacy",
    "ineffective",
    "primary end point was not met",
    "primary endpoint was not met",
    "limited benefit",
    "did not prolong",
    "did not improve survival",
    "not superior"
]

OPPOSITE_TITLE_HINTS = [
    "worse",
    "harm",
    "adverse",
    "toxicity",
    "increased risk",
    "negative effect",
    "inferior"
]


def infer_pubmed_result_type(article):
    """
    Very rough text heuristic for prototype labeling.
    """
    text = " ".join([
        article.get("title", ""),
        article.get("abstract", ""),
        " ".join(article.get("mesh_terms", []) or [])
    ]).lower()

    if any(term in text for term in OPPOSITE_TITLE_HINTS):
        return "opposite"
    if any(term in text for term in NEGATIVE_TITLE_HINTS):
        return "null"
    return "inconclusive"


def fallback_graph_tags_for_trial(trial):
    conditions = trial.get("conditions", []) or []
    interventions = trial.get("interventions", []) or []

    return {
        "mechanism": "unspecified",
        "target": interventions[0] if interventions else "unspecified",
        "population": trial.get("population", "") or "unspecified",
        "therapeutic_area": conditions[0] if conditions else "unspecified"
    }


def fallback_graph_tags_for_pubmed(article):
    mesh_terms = article.get("mesh_terms", []) or []
    journal = article.get("journal", "") or ""

    therapeutic_area = mesh_terms[0] if mesh_terms else (journal or "unspecified")
    target = mesh_terms[1] if len(mesh_terms) > 1 else "unspecified"

    return {
        "mechanism": "literature-derived",
        "target": target,
        "population": "unspecified",
        "therapeutic_area": therapeutic_area or "unspecified"
    }


# ---------------------------
# Claude JSON parsing helpers
# ---------------------------

def clean_model_json_text(raw_text):
    text = (raw_text or "").strip()

    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)

    return text.strip()


def safe_parse_json(raw_text):
    text = clean_model_json_text(raw_text)

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        text = text.replace("\r", " ")
        text = re.sub(r"[\x00-\x1f]+", " ", text)
        text = re.sub(r"\s+", " ", text).strip()
        return json.loads(text)


# ---------------------------
# Claude summarization
# ---------------------------

def summarize_with_claude(
    title,
    study_type,
    result_type,
    intervention,
    population,
    primary_outcome,
    p_value,
    effect_estimate,
    limitations
):
    short_limitations = (limitations or "")[:4000]

    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=500,
        system=(
            "You are a scientific summarization assistant for FailUp. "
            "Return only valid raw JSON, with no markdown fences or extra text. "
            "Be precise, neutral, concise, and cautious. "
            "Do not speculate beyond what was provided. "
            "If information is insufficient, say so plainly."
        ),
        messages=[
            {
                "role": "user",
                "content": f"""
Return ONLY valid JSON with these keys:
background
findings
main_limitation
failure_mode
contradiction_check
graph_tags

Rules:
- Return raw JSON only. No markdown. No backticks.
- Keep each field concise.
- failure_mode must be one of:
  "target validity"
  "patient selection"
  "dosing/PK"
  "underpowered"
  "outcome measurement"
  "off-target effects"
  "unknown"
- contradiction_check should be "Not assessed from provided information."
- graph_tags must contain:
  mechanism
  target
  population
  therapeutic_area

Study title: {title}
Study type: {study_type}
Result type: {result_type}
Intervention: {intervention}
Population: {population}
Primary outcome: {primary_outcome}
P-value: {p_value}
Effect estimate: {effect_estimate}
Limitations or context: {short_limitations}
"""
            }
        ]
    )

    raw_text = response.content[0].text
    return safe_parse_json(raw_text)


# ---------------------------
# Claude curation
# ---------------------------

def curate_failure_with_claude(
    title,
    study_type,
    result_type_hint,
    source_type,
    intervention,
    population,
    primary_outcome,
    abstract_or_summary,
    extra_context
):
    short_text = (abstract_or_summary or "")[:5000]
    short_context = (extra_context or "")[:2500]

    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=400,
        system=(
            "You are a scientific screening assistant for FailUp. "
            "Your task is to decide whether a record is a genuine failure-related study record suitable for the FailUp repository. "
            "Return only valid raw JSON. No markdown. No extra words. "
            "Be conservative. If the evidence is weak, choose uncertain."
        ),
        messages=[
            {
                "role": "user",
                "content": f"""
Return ONLY valid JSON with these keys:
review_status
confidence
basis
graph_worthy
failure_class

Rules:
- review_status must be one of:
  "curated_failure"
  "uncertain"
  "not_failure"
- confidence must be one of:
  "high"
  "medium"
  "low"
- graph_worthy must be true or false
- failure_class must be one of:
  "null"
  "inconclusive"
  "opposite"
  "not_failure"
- Be conservative.
- Only choose "curated_failure" if the provided material supports that this is a failure-type record:
  failed / null / negative / inconclusive / not meeting primary endpoint / opposite effect / trial stopped for a failure-related reason.
- Do NOT assume that completed means failed.
- If the record is a positive or clearly successful study, return "not_failure".
- If evidence is too limited, return "uncertain".
- basis should be 1-2 concise sentences explaining why.

Title: {title}
Source type: {source_type}
Study type: {study_type}
Current result type hint: {result_type_hint}
Intervention: {intervention}
Population: {population}
Primary outcome: {primary_outcome}
Abstract or summary: {short_text}
Extra context: {short_context}
"""
            }
        ]
    )

    raw_text = response.content[0].text
    data = safe_parse_json(raw_text)

    review_status = data.get("review_status", "uncertain")
    if review_status not in {"curated_failure", "uncertain", "not_failure"}:
        review_status = "uncertain"

    confidence = data.get("confidence", "low")
    if confidence not in {"high", "medium", "low"}:
        confidence = "low"

    failure_class = data.get("failure_class", "not_failure")
    if failure_class not in {"null", "inconclusive", "opposite", "not_failure"}:
        failure_class = "not_failure"

    graph_worthy = bool(data.get("graph_worthy", False))
    basis = data.get("basis", "No basis returned.")

    return {
        "review_status": review_status,
        "confidence": confidence,
        "basis": basis,
        "graph_worthy": graph_worthy,
        "failure_class": failure_class,
        "reviewed_at": now_iso()
    }


# ---------------------------
# Graph helpers (kept for backend compatibility)
# ---------------------------

def normalize_result_type(value):
    value = normalize_text(value)
    if value in {"null", "inconclusive", "opposite"}:
        return value
    return "inconclusive"


def standardize_graph_tags(tags):
    tags = tags or {}
    return {
        "mechanism": sanitize_graph_value(tags.get("mechanism", "unspecified")),
        "target": sanitize_graph_value(tags.get("target", "unspecified")),
        "population": sanitize_graph_value(tags.get("population", "unspecified")),
        "therapeutic_area": sanitize_graph_value(tags.get("therapeutic_area", "unspecified")),
    }


def contradiction_flag(value):
    text = normalize_text(value)
    return bool(text and text != "not assessed from provided information.")


def is_graph_eligible_imported(record):
    curation = record.get("curation", {}) or {}
    return (
        curation.get("review_status") == "curated_failure" and
        curation.get("graph_worthy") is True
    )


def record_to_study_node(record):
    tags = standardize_graph_tags(record.get("graph_tags", {}))
    return {
        "id": record["id"],
        "kind": "study",
        "label": truncate_label(record["title"], 24),
        "title": record["title"],
        "type": normalize_result_type(record.get("type")),
        "doi": record.get("doi", ""),
        "failure_mode": record.get("failure_mode", "unknown"),
        "contradiction_check": record.get("contradiction_check", ""),
        "has_contradiction": contradiction_flag(record.get("contradiction_check", "")),
        "source": record.get("source", ""),
        "source_label": record.get("source_label", ""),
        "graph_tags": tags
    }


def collect_graph_records(submissions, imported_trials=None, pubmed_articles=None):
    all_records = []

    # Manual submissions always eligible
    for sub in submissions or []:
        all_records.append({
            "id": f"manual-{sub.get('id')}",
            "title": safe_title(sub.get("title"), "Untitled submission"),
            "type": normalize_result_type(sub.get("result_type", "null")),
            "doi": sub.get("doi", ""),
            "failure_mode": sub.get("summary", {}).get("failure_mode", "unknown"),
            "contradiction_check": sub.get("summary", {}).get("contradiction_check", ""),
            "graph_tags": sub.get("summary", {}).get("graph_tags", {}),
            "source": "manual",
            "source_label": SOURCE_LABELS["manual"]
        })

    for trial in imported_trials or []:
        if not is_graph_eligible_imported(trial):
            continue

        graph_tags = trial.get("summary", {}).get("graph_tags") or fallback_graph_tags_for_trial(trial)
        failure_class = trial.get("curation", {}).get("failure_class") or trial.get("result_type") or infer_trial_result_type(trial)

        all_records.append({
            "id": f"ctgov-{trial.get('id')}",
            "title": safe_title(trial.get("title"), "Imported trial"),
            "type": normalize_result_type(failure_class),
            "doi": trial.get("nct_id", ""),
            "failure_mode": trial.get("summary", {}).get("failure_mode", "unknown"),
            "contradiction_check": trial.get("summary", {}).get("contradiction_check", ""),
            "graph_tags": graph_tags,
            "source": "ctgov_curated",
            "source_label": SOURCE_LABELS["ctgov_curated"]
        })

    for article in pubmed_articles or []:
        if not is_graph_eligible_imported(article):
            continue

        graph_tags = article.get("summary", {}).get("graph_tags") or fallback_graph_tags_for_pubmed(article)
        failure_class = article.get("curation", {}).get("failure_class") or article.get("result_type") or infer_pubmed_result_type(article)

        all_records.append({
            "id": f"pubmed-{article.get('id')}",
            "title": safe_title(article.get("title"), "Imported article"),
            "type": normalize_result_type(failure_class),
            "doi": article.get("pmid", ""),
            "failure_mode": article.get("summary", {}).get("failure_mode", "unknown"),
            "contradiction_check": article.get("summary", {}).get("contradiction_check", ""),
            "graph_tags": graph_tags,
            "source": "pubmed_curated",
            "source_label": SOURCE_LABELS["pubmed_curated"]
        })

    return all_records


def build_study_graph(all_records):
    nodes = [record_to_study_node(record) for record in all_records]
    edges = []
    degree_map = {node["id"]: 0 for node in nodes}

    for i in range(len(nodes)):
        for j in range(i + 1, len(nodes)):
            n1 = nodes[i]
            n2 = nodes[j]

            shared = []
            for field in ["mechanism", "target", "population", "therapeutic_area"]:
                v1 = normalize_text(n1["graph_tags"].get(field, ""))
                v2 = normalize_text(n2["graph_tags"].get(field, ""))

                if v1 and v2 and v1 == v2 and v1 != "unspecified":
                    shared.append(field)

            strong_link = (
                "target" in shared or
                "mechanism" in shared or
                len(shared) >= 2
            )

            if strong_link:
                edge = {
                    "source": n1["id"],
                    "target": n2["id"],
                    "shared_tags": shared,
                    "weight": len(shared),
                    "kind": "study-study"
                }
                edges.append(edge)
                degree_map[n1["id"]] += 1
                degree_map[n2["id"]] += 1

    for node in nodes:
        node["degree"] = degree_map.get(node["id"], 0)

    return {"nodes": nodes, "edges": edges}


def build_overview_graph(study_nodes):
    concept_nodes = {}
    edge_weights = {}

    def add_concept(kind, raw_value, study_node):
        value = sanitize_graph_value(raw_value, "unspecified")
        if normalize_text(value) == "unspecified":
            return None

        concept_id = f"{kind}:{normalize_text(value)}"
        if concept_id not in concept_nodes:
            concept_nodes[concept_id] = {
                "id": concept_id,
                "kind": kind,
                "label": truncate_label(value, 28),
                "title": value,
                "type": "concept",
                "count": 0,
                "sources": set(),
                "result_type_counts": {"null": 0, "inconclusive": 0, "opposite": 0},
                "study_ids": []
            }

        concept_nodes[concept_id]["count"] += 1
        concept_nodes[concept_id]["sources"].add(study_node["source"])
        concept_nodes[concept_id]["study_ids"].append(study_node["id"])
        concept_nodes[concept_id]["result_type_counts"][study_node["type"]] += 1
        return concept_id

    for node in study_nodes:
        area_id = add_concept("therapeutic_area", node["graph_tags"]["therapeutic_area"], node)
        target_id = add_concept("target", node["graph_tags"]["target"], node)
        mechanism_id = add_concept("mechanism", node["graph_tags"]["mechanism"], node)

        if area_id and target_id:
            edge_weights[(area_id, target_id)] = edge_weights.get((area_id, target_id), 0) + 1
        if target_id and mechanism_id:
            edge_weights[(target_id, mechanism_id)] = edge_weights.get((target_id, mechanism_id), 0) + 1

    overview_nodes = []
    for node in concept_nodes.values():
        node["sources"] = sorted(list(node["sources"]))
        overview_nodes.append(node)

    overview_edges = [
        {
            "source": source_id,
            "target": target_id,
            "weight": weight,
            "kind": "concept-concept"
        }
        for (source_id, target_id), weight in edge_weights.items()
    ]

    return {"nodes": overview_nodes, "edges": overview_edges}


def build_graph_data(submissions, imported_trials=None, pubmed_articles=None):
    all_records = collect_graph_records(submissions, imported_trials, pubmed_articles)
    studies = build_study_graph(all_records)
    overview = build_overview_graph(studies["nodes"])

    return {
        "nodes": studies["nodes"],
        "edges": studies["edges"],
        "studies": studies,
        "overview": overview,
        "meta": {
            "default_view": "overview",
            "study_node_count": len(studies["nodes"]),
            "study_edge_count": len(studies["edges"]),
            "overview_node_count": len(overview["nodes"]),
            "overview_edge_count": len(overview["edges"])
        }
    }


# ---------------------------
# CT.gov helpers
# ---------------------------

def fetch_ctgov_page(query, page_size, page_token=None):
    params = {
        "query.term": query,
        "pageSize": page_size
    }
    if page_token:
        params["pageToken"] = page_token

    response = requests.get(CTGOV_API_URL, params=params, timeout=30)
    response.raise_for_status()
    payload = response.json()

    return payload.get("studies", []), payload.get("nextPageToken")


# ---------------------------
# PubMed helpers
# ---------------------------

def ncbi_common_params():
    params = {"tool": NCBI_TOOL}
    if NCBI_EMAIL:
        params["email"] = NCBI_EMAIL
    if NCBI_API_KEY:
        params["api_key"] = NCBI_API_KEY
    return params


def pubmed_esearch(query, retmax=20, retstart=0):
    params = {
        **ncbi_common_params(),
        "db": "pubmed",
        "term": query,
        "retmax": retmax,
        "retstart": retstart,
        "retmode": "json",
        "sort": "relevance"
    }
    response = requests.get(f"{NCBI_EUTILS_BASE}/esearch.fcgi", params=params, timeout=30)
    response.raise_for_status()
    return response.json()


def pubmed_esummary(pmids):
    if not pmids:
        return {}
    params = {
        **ncbi_common_params(),
        "db": "pubmed",
        "id": ",".join(pmids),
        "retmode": "json"
    }
    response = requests.get(f"{NCBI_EUTILS_BASE}/esummary.fcgi", params=params, timeout=30)
    response.raise_for_status()
    return response.json()


def pubmed_efetch_xml(pmids):
    if not pmids:
        return ""
    params = {
        **ncbi_common_params(),
        "db": "pubmed",
        "id": ",".join(pmids),
        "retmode": "xml"
    }
    response = requests.get(f"{NCBI_EUTILS_BASE}/efetch.fcgi", params=params, timeout=30)
    response.raise_for_status()
    return response.text


def parse_pubmed_xml_abstracts_and_mesh(xml_text):
    abstracts_by_pmid = {}
    mesh_by_pmid = {}

    if not xml_text.strip():
        return abstracts_by_pmid, mesh_by_pmid

    root = ET.fromstring(xml_text)

    for article in root.findall(".//PubmedArticle"):
        pmid_el = article.find(".//MedlineCitation/PMID")
        if pmid_el is None or not pmid_el.text:
            continue
        pmid = pmid_el.text.strip()

        abstract_parts = []
        for abst in article.findall(".//Abstract/AbstractText"):
            label = abst.attrib.get("Label", "").strip()
            text = "".join(abst.itertext()).strip()
            if text:
                abstract_parts.append(f"{label}: {text}" if label else text)
        abstracts_by_pmid[pmid] = "\n".join(abstract_parts).strip()

        mesh_terms = []
        for mh in article.findall(".//MeshHeading/DescriptorName"):
            txt = "".join(mh.itertext()).strip()
            if txt:
                mesh_terms.append(txt)
        mesh_by_pmid[pmid] = mesh_terms

    return abstracts_by_pmid, mesh_by_pmid


# ---------------------------
# Routes
# ---------------------------

@app.route("/")
def home():
    return send_file("failup.html")


@app.route("/summarize", methods=["POST"])
def summarize():
    data = request.get_json() or {}

    try:
        summary = summarize_with_claude(
            title=data.get("title", ""),
            study_type=data.get("study_type", ""),
            result_type=data.get("result_type", ""),
            intervention=data.get("intervention", ""),
            population=data.get("population", ""),
            primary_outcome=data.get("primary_outcome", ""),
            p_value=data.get("p_value", ""),
            effect_estimate=data.get("effect_estimate", ""),
            limitations=data.get("limitations", "")
        )

        submissions = load_submissions()
        doi = generate_doi(len(submissions) + 1)

        record = {
            "id": len(submissions) + 1,
            "doi": doi,
            "submitted_at": now_iso(),
            "title": data.get("title", ""),
            "study_type": data.get("study_type", ""),
            "result_type": data.get("result_type", ""),
            "intervention": data.get("intervention", ""),
            "population": data.get("population", ""),
            "primary_outcome": data.get("primary_outcome", ""),
            "p_value": data.get("p_value", ""),
            "effect_estimate": data.get("effect_estimate", ""),
            "limitations": data.get("limitations", ""),
            "summary": summary
        }

        submissions.insert(0, record)
        save_submissions(submissions)

        return jsonify({
            "success": True,
            "doi": doi,
            "summary": summary,
            "record": record
        })

    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route("/extract-study-fields", methods=["POST"])
def extract_study_fields():
    files = request.files.getlist("files")
    files = [f for f in files if f and (f.filename or "").strip()]

    if not files:
        return jsonify({
            "success": False,
            "error": "No files were uploaded."
        }), 400

    try:
        document_text = build_uploaded_document_context(files)
        if not document_text.strip():
            return jsonify({
                "success": False,
                "error": "The uploaded files did not contain extractable text."
            }), 400

        extracted_fields = extract_study_fields_with_claude(
            document_text=document_text,
            file_names=[f.filename or "uploaded_file" for f in files]
        )

        return jsonify({
            "success": True,
            "fields": extracted_fields,
            "file_names": [f.filename or "uploaded_file" for f in files]
        })

    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route("/submissions", methods=["GET"])
def get_submissions():
    return jsonify({
        "success": True,
        "submissions": load_submissions()
    })


@app.route("/graph-data", methods=["GET"])
def get_graph_data():
    submissions = load_submissions()
    trials = load_trials()
    pubmed_articles = load_pubmed_articles()

    graph_data = build_graph_data(submissions, trials, pubmed_articles)

    return jsonify({
        "success": True,
        "graph": graph_data
    })


# ---------------------------
# CT.gov routes
# ---------------------------

@app.route("/ingest-ctgov", methods=["POST"])
def ingest_ctgov():
    data = request.get_json() or {}
    query = (data.get("query") or "").strip()
    page_size = safe_int(data.get("page_size", 20), 20)
    max_pages = safe_int(data.get("max_pages", 5), 5)

    page_size = max(1, min(page_size, 100))
    max_pages = max(1, min(max_pages, 50))

    if not query:
        return jsonify({
            "success": False,
            "error": "Missing query."
        }), 400

    try:
        existing_trials = load_trials()
        existing_nct_ids = {trial.get("nct_id") for trial in existing_trials if trial.get("nct_id")}

        new_records = []
        page_token = None
        pages_fetched = 0
        fetched_count = 0
        skipped_duplicates = 0

        while pages_fetched < max_pages:
            studies, next_page_token = fetch_ctgov_page(
                query=query,
                page_size=page_size,
                page_token=page_token
            )
            pages_fetched += 1

            if not studies:
                break

            fetched_count += len(studies)

            for study in studies:
                protocol = study.get("protocolSection", {})
                ident = protocol.get("identificationModule", {})
                status = protocol.get("statusModule", {})
                design = protocol.get("designModule", {})
                conditions_mod = protocol.get("conditionsModule", {})
                arms_mod = protocol.get("armsInterventionsModule", {})
                eligibility = protocol.get("eligibilityModule", {})
                outcomes = protocol.get("outcomesModule", {})
                description = protocol.get("descriptionModule", {})

                nct_id = ident.get("nctId", "")
                if not nct_id:
                    continue

                if nct_id in existing_nct_ids:
                    skipped_duplicates += 1
                    continue

                title = ident.get("briefTitle") or ident.get("officialTitle") or nct_id
                phases = design.get("phases", [])

                interventions = []
                for item in arms_mod.get("interventions", []):
                    name = item.get("name", "")
                    itype = item.get("type", "")
                    if name:
                        interventions.append(f"{itype}: {name}" if itype else name)

                primary_outcomes = outcomes.get("primaryOutcomes", [])
                primary_outcome = ""
                if primary_outcomes:
                    primary_outcome = primary_outcomes[0].get("measure", "")

                population_parts = []
                sex = eligibility.get("sex", "")
                min_age = eligibility.get("minimumAge", "")
                max_age = eligibility.get("maximumAge", "")
                healthy = eligibility.get("healthyVolunteers", "")
                if sex:
                    population_parts.append(f"Sex: {sex}")
                if min_age or max_age:
                    population_parts.append(f"Age: {min_age} to {max_age}".strip())
                if healthy:
                    population_parts.append(f"Healthy volunteers: {healthy}")

                record = {
                    "id": len(existing_trials) + len(new_records) + 1,
                    "source": "clinicaltrials.gov",
                    "source_status": "ctgov_imported",
                    "nct_id": nct_id,
                    "title": title,
                    "study_type": design.get("studyType", ""),
                    "phase": ", ".join(phases) if phases else "",
                    "status": status.get("overallStatus", ""),
                    "conditions": conditions_mod.get("conditions", []),
                    "interventions": interventions,
                    "population": " | ".join([p for p in population_parts if p]),
                    "primary_outcome": primary_outcome,
                    "brief_summary": description.get("briefSummary", ""),
                    "url": f"https://clinicaltrials.gov/study/{nct_id}",
                    "imported_at": now_iso(),
                    "result_type": infer_trial_result_type({"status": status.get("overallStatus", "")})
                }

                new_records.append(record)
                existing_nct_ids.add(nct_id)

            if not next_page_token:
                break

            page_token = next_page_token

        if new_records:
            existing_trials = new_records + existing_trials
            save_trials(existing_trials)

        return jsonify({
            "success": True,
            "query": query,
            "page_size": page_size,
            "max_pages": max_pages,
            "pages_fetched": pages_fetched,
            "fetched_count": fetched_count,
            "imported_count": len(new_records),
            "skipped_duplicates": skipped_duplicates,
            "total_trials": len(existing_trials),
            "records": new_records
        })

    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route("/ctgov-trials", methods=["GET"])
def get_ctgov_trials():
    trials = load_trials()
    imported_count = sum(1 for t in trials if classify_trial_source(t) == "ctgov_imported")
    enriched_count = sum(1 for t in trials if classify_trial_source(t) == "ctgov_enriched")
    curated_count = sum(1 for t in trials if classify_trial_source(t) == "ctgov_curated")

    return jsonify({
        "success": True,
        "trials": trials,
        "counts": {
            "total": len(trials),
            "ctgov_imported": imported_count,
            "ctgov_enriched": enriched_count,
            "ctgov_curated": curated_count
        },
        "label_map": {
            "ctgov_imported": SOURCE_LABELS["ctgov_imported"],
            "ctgov_enriched": SOURCE_LABELS["ctgov_enriched"],
            "ctgov_curated": SOURCE_LABELS["ctgov_curated"]
        }
    })


@app.route("/enrich-ctgov-trials", methods=["POST"])
def enrich_ctgov_trials():
    data = request.get_json() or {}
    limit = safe_int(data.get("limit", 1), 1)

    try:
        trials = load_trials()
        updated_count = 0
        already_enriched = 0
        errors = []

        for trial in trials:
            if updated_count >= limit:
                break

            if "summary" in trial:
                already_enriched += 1
                continue

            title = trial.get("title", "")
            result_type = infer_trial_result_type(trial).capitalize()
            intervention_text = "; ".join(trial.get("interventions", []))
            condition_text = "; ".join(trial.get("conditions", []))
            context_text = (
                f"{trial.get('brief_summary', '')[:2500]}\n"
                f"Conditions: {condition_text}\n"
                f"Status: {trial.get('status', '')}\n"
                f"Phase: {trial.get('phase', '')}"
            )

            try:
                summary = summarize_with_claude(
                    title=title,
                    study_type=trial.get("study_type", ""),
                    result_type=result_type,
                    intervention=intervention_text,
                    population=trial.get("population", ""),
                    primary_outcome=trial.get("primary_outcome", ""),
                    p_value="",
                    effect_estimate="",
                    limitations=context_text
                )

                trial["result_type"] = normalize_text(result_type)
                trial["summary"] = summary
                trial["source_status"] = "ctgov_enriched"
                trial["enriched_at"] = now_iso()

                updated_count += 1
                save_trials(trials)

            except Exception as trial_error:
                errors.append({
                    "title": title,
                    "error": str(trial_error)
                })

        imported_count = sum(1 for t in trials if classify_trial_source(t) == "ctgov_imported")
        enriched_count = sum(1 for t in trials if classify_trial_source(t) == "ctgov_enriched")
        curated_count = sum(1 for t in trials if classify_trial_source(t) == "ctgov_curated")

        return jsonify({
            "success": True,
            "enriched_count": updated_count,
            "already_enriched_seen": already_enriched,
            "total_trials": len(trials),
            "counts": {
                "ctgov_imported": imported_count,
                "ctgov_enriched": enriched_count,
                "ctgov_curated": curated_count
            },
            "label_map": {
                "ctgov_imported": SOURCE_LABELS["ctgov_imported"],
                "ctgov_enriched": SOURCE_LABELS["ctgov_enriched"],
                "ctgov_curated": SOURCE_LABELS["ctgov_curated"]
            },
            "errors": errors
        })

    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route("/curate-ctgov-trials", methods=["POST"])
def curate_ctgov_trials():
    data = request.get_json() or {}
    limit = safe_int(data.get("limit", 5), 5)

    try:
        trials = load_trials()
        curated_count = 0
        already_curated_seen = 0
        errors = []

        for trial in trials:
            if curated_count >= limit:
                break

            if trial.get("curation"):
                already_curated_seen += 1
                continue

            summary = trial.get("summary", {}) or {}
            intervention_text = "; ".join(trial.get("interventions", []))
            context_text = (
                f"Status: {trial.get('status', '')}\n"
                f"Phase: {trial.get('phase', '')}\n"
                f"Conditions: {'; '.join(trial.get('conditions', []))}\n"
                f"Summary findings: {summary.get('findings', '')}\n"
                f"Main limitation: {summary.get('main_limitation', '')}\n"
                f"Failure mode: {summary.get('failure_mode', '')}"
            )

            try:
                curation = curate_failure_with_claude(
                    title=trial.get("title", ""),
                    study_type=trial.get("study_type", ""),
                    result_type_hint=trial.get("result_type", infer_trial_result_type(trial)),
                    source_type="ClinicalTrials.gov record",
                    intervention=intervention_text,
                    population=trial.get("population", ""),
                    primary_outcome=trial.get("primary_outcome", ""),
                    abstract_or_summary=trial.get("brief_summary", ""),
                    extra_context=context_text
                )

                trial["curation"] = curation
                if curation.get("review_status") == "curated_failure" and curation.get("failure_class") in {"null", "inconclusive", "opposite"}:
                    trial["result_type"] = curation["failure_class"]

                if curation.get("review_status") == "curated_failure" and curation.get("graph_worthy") is True:
                    trial["source_status"] = "ctgov_curated"

                curated_count += 1
                save_trials(trials)

            except Exception as trial_error:
                errors.append({
                    "title": trial.get("title", "Unknown title"),
                    "error": str(trial_error)
                })

        imported_count = sum(1 for t in trials if classify_trial_source(t) == "ctgov_imported")
        enriched_count = sum(1 for t in trials if classify_trial_source(t) == "ctgov_enriched")
        curated_total = sum(1 for t in trials if classify_trial_source(t) == "ctgov_curated")

        return jsonify({
            "success": True,
            "curated_count": curated_count,
            "already_curated_seen": already_curated_seen,
            "total_trials": len(trials),
            "counts": {
                "ctgov_imported": imported_count,
                "ctgov_enriched": enriched_count,
                "ctgov_curated": curated_total
            },
            "label_map": {
                "ctgov_imported": SOURCE_LABELS["ctgov_imported"],
                "ctgov_enriched": SOURCE_LABELS["ctgov_enriched"],
                "ctgov_curated": SOURCE_LABELS["ctgov_curated"]
            },
            "errors": errors
        })

    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


# ---------------------------
# PubMed routes
# ---------------------------

@app.route("/ingest-pubmed", methods=["POST"])
def ingest_pubmed():
    data = request.get_json() or {}
    query = (data.get("query") or "").strip()
    retmax = safe_int(data.get("retmax", 20), 20)
    retstart = safe_int(data.get("retstart", 0), 0)

    retmax = max(1, min(retmax, 200))
    retstart = max(0, retstart)

    if not query:
        return jsonify({
            "success": False,
            "error": "Missing query."
        }), 400

    try:
        existing_articles = load_pubmed_articles()
        existing_pmids = {a.get("pmid") for a in existing_articles if a.get("pmid")}

        esearch_payload = pubmed_esearch(query=query, retmax=retmax, retstart=retstart)
        idlist = esearch_payload.get("esearchresult", {}).get("idlist", []) or []
        total_count = safe_int(esearch_payload.get("esearchresult", {}).get("count", 0), 0)

        if not idlist:
            return jsonify({
                "success": True,
                "query": query,
                "retmax": retmax,
                "retstart": retstart,
                "fetched_count": 0,
                "imported_count": 0,
                "skipped_duplicates": 0,
                "total_pubmed": len(existing_articles),
                "records": [],
                "pubmed_total_matches": total_count
            })

        summary_payload = pubmed_esummary(idlist)
        summary_result = summary_payload.get("result", {}) or {}

        xml_text = pubmed_efetch_xml(idlist)
        abstracts_by_pmid, mesh_by_pmid = parse_pubmed_xml_abstracts_and_mesh(xml_text)

        new_records = []
        skipped_duplicates = 0

        for pmid in idlist:
            if pmid in existing_pmids:
                skipped_duplicates += 1
                continue

            item = summary_result.get(pmid, {}) or {}
            title = item.get("title", "")
            article_ids = item.get("articleids", []) or []

            doi = ""
            for aid in article_ids:
                if aid.get("idtype") == "doi":
                    doi = aid.get("value", "")
                    break

            authors = []
            for a in item.get("authors", []) or []:
                if a.get("name"):
                    authors.append(a["name"])

            article = {
                "id": len(existing_articles) + len(new_records) + 1,
                "source": "pubmed",
                "source_status": "pubmed_imported",
                "pmid": pmid,
                "doi": doi,
                "title": title,
                "journal": item.get("fulljournalname") or item.get("source", ""),
                "pub_date": item.get("pubdate", ""),
                "authors": authors,
                "article_type": ", ".join(item.get("pubtype", []) or []),
                "abstract": abstracts_by_pmid.get(pmid, ""),
                "mesh_terms": mesh_by_pmid.get(pmid, []),
                "url": f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/",
                "imported_at": now_iso()
            }

            article["result_type"] = infer_pubmed_result_type(article)

            new_records.append(article)
            existing_pmids.add(pmid)

        if new_records:
            existing_articles = new_records + existing_articles
            save_pubmed_articles(existing_articles)

        return jsonify({
            "success": True,
            "query": query,
            "retmax": retmax,
            "retstart": retstart,
            "fetched_count": len(idlist),
            "imported_count": len(new_records),
            "skipped_duplicates": skipped_duplicates,
            "total_pubmed": len(existing_articles),
            "records": new_records,
            "pubmed_total_matches": total_count
        })

    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route("/pubmed-articles", methods=["GET"])
def get_pubmed_articles():
    articles = load_pubmed_articles()
    imported_count = sum(1 for a in articles if classify_pubmed_source(a) == "pubmed_imported")
    enriched_count = sum(1 for a in articles if classify_pubmed_source(a) == "pubmed_enriched")
    curated_count = sum(1 for a in articles if classify_pubmed_source(a) == "pubmed_curated")

    return jsonify({
        "success": True,
        "articles": articles,
        "counts": {
            "total": len(articles),
            "pubmed_imported": imported_count,
            "pubmed_enriched": enriched_count,
            "pubmed_curated": curated_count
        },
        "label_map": {
            "pubmed_imported": SOURCE_LABELS["pubmed_imported"],
            "pubmed_enriched": SOURCE_LABELS["pubmed_enriched"],
            "pubmed_curated": SOURCE_LABELS["pubmed_curated"]
        }
    })


@app.route("/enrich-pubmed-articles", methods=["POST"])
def enrich_pubmed_articles():
    data = request.get_json() or {}
    limit = safe_int(data.get("limit", 1), 1)

    try:
        articles = load_pubmed_articles()
        updated_count = 0
        already_enriched = 0
        errors = []

        for article in articles:
            if updated_count >= limit:
                break

            if "summary" in article:
                already_enriched += 1
                continue

            title = article.get("title", "")
            result_type = normalize_text(article.get("result_type", infer_pubmed_result_type(article))).capitalize()
            abstract_short = (article.get("abstract", "") or "")[:2500]

            limitations = (
                f"Journal: {article.get('journal', '')}\n"
                f"Publication date: {article.get('pub_date', '')}\n"
                f"Article type: {article.get('article_type', '')}\n"
                f"Abstract: {abstract_short}\n"
                f"MeSH terms: {', '.join(article.get('mesh_terms', []) or [])}"
            )

            try:
                summary = summarize_with_claude(
                    title=title,
                    study_type=article.get("article_type", "PubMed article"),
                    result_type=result_type,
                    intervention="Not explicitly structured in PubMed metadata.",
                    population="Not explicitly structured in PubMed metadata.",
                    primary_outcome="Not explicitly structured in PubMed metadata.",
                    p_value="",
                    effect_estimate="",
                    limitations=limitations
                )

                article["result_type"] = normalize_text(result_type)
                article["summary"] = summary
                article["source_status"] = "pubmed_enriched"
                article["enriched_at"] = now_iso()

                updated_count += 1
                save_pubmed_articles(articles)

            except Exception as article_error:
                errors.append({
                    "title": title,
                    "error": str(article_error)
                })

        imported_count = sum(1 for a in articles if classify_pubmed_source(a) == "pubmed_imported")
        enriched_count = sum(1 for a in articles if classify_pubmed_source(a) == "pubmed_enriched")
        curated_count = sum(1 for a in articles if classify_pubmed_source(a) == "pubmed_curated")

        return jsonify({
            "success": True,
            "enriched_count": updated_count,
            "already_enriched_seen": already_enriched,
            "total_articles": len(articles),
            "counts": {
                "pubmed_imported": imported_count,
                "pubmed_enriched": enriched_count,
                "pubmed_curated": curated_count
            },
            "label_map": {
                "pubmed_imported": SOURCE_LABELS["pubmed_imported"],
                "pubmed_enriched": SOURCE_LABELS["pubmed_enriched"],
                "pubmed_curated": SOURCE_LABELS["pubmed_curated"]
            },
            "errors": errors
        })

    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route("/curate-pubmed-articles", methods=["POST"])
def curate_pubmed_articles():
    data = request.get_json() or {}
    limit = safe_int(data.get("limit", 5), 5)

    try:
        articles = load_pubmed_articles()
        curated_count = 0
        already_curated_seen = 0
        errors = []

        for article in articles:
            if curated_count >= limit:
                break

            if article.get("curation"):
                already_curated_seen += 1
                continue

            summary = article.get("summary", {}) or {}
            context_text = (
                f"Journal: {article.get('journal', '')}\n"
                f"Publication date: {article.get('pub_date', '')}\n"
                f"Article type: {article.get('article_type', '')}\n"
                f"Summary findings: {summary.get('findings', '')}\n"
                f"Main limitation: {summary.get('main_limitation', '')}\n"
                f"Failure mode: {summary.get('failure_mode', '')}\n"
                f"MeSH terms: {', '.join(article.get('mesh_terms', []) or [])}"
            )

            try:
                curation = curate_failure_with_claude(
                    title=article.get("title", ""),
                    study_type=article.get("article_type", ""),
                    result_type_hint=article.get("result_type", infer_pubmed_result_type(article)),
                    source_type="PubMed record",
                    intervention="Not explicitly structured in PubMed metadata.",
                    population="Not explicitly structured in PubMed metadata.",
                    primary_outcome="Not explicitly structured in PubMed metadata.",
                    abstract_or_summary=article.get("abstract", ""),
                    extra_context=context_text
                )

                article["curation"] = curation
                if curation.get("review_status") == "curated_failure" and curation.get("failure_class") in {"null", "inconclusive", "opposite"}:
                    article["result_type"] = curation["failure_class"]

                if curation.get("review_status") == "curated_failure" and curation.get("graph_worthy") is True:
                    article["source_status"] = "pubmed_curated"

                curated_count += 1
                save_pubmed_articles(articles)

            except Exception as article_error:
                errors.append({
                    "title": article.get("title", "Unknown title"),
                    "error": str(article_error)
                })

        imported_count = sum(1 for a in articles if classify_pubmed_source(a) == "pubmed_imported")
        enriched_count = sum(1 for a in articles if classify_pubmed_source(a) == "pubmed_enriched")
        curated_total = sum(1 for a in articles if classify_pubmed_source(a) == "pubmed_curated")

        return jsonify({
            "success": True,
            "curated_count": curated_count,
            "already_curated_seen": already_curated_seen,
            "total_articles": len(articles),
            "counts": {
                "pubmed_imported": imported_count,
                "pubmed_enriched": enriched_count,
                "pubmed_curated": curated_total
            },
            "label_map": {
                "pubmed_imported": SOURCE_LABELS["pubmed_imported"],
                "pubmed_enriched": SOURCE_LABELS["pubmed_enriched"],
                "pubmed_curated": SOURCE_LABELS["pubmed_curated"]
            },
            "errors": errors
        })

    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


# ---------------------------
# Combined source counts
# ---------------------------

@app.route("/source-counts", methods=["GET"])
def source_counts():
    submissions = load_submissions()
    trials = load_trials()
    articles = load_pubmed_articles()

    ctgov_imported = sum(1 for t in trials if classify_trial_source(t) == "ctgov_imported")
    ctgov_enriched = sum(1 for t in trials if classify_trial_source(t) == "ctgov_enriched")
    ctgov_curated = sum(1 for t in trials if classify_trial_source(t) == "ctgov_curated")

    pubmed_imported = sum(1 for a in articles if classify_pubmed_source(a) == "pubmed_imported")
    pubmed_enriched = sum(1 for a in articles if classify_pubmed_source(a) == "pubmed_enriched")
    pubmed_curated = sum(1 for a in articles if classify_pubmed_source(a) == "pubmed_curated")

    return jsonify({
        "success": True,
        "counts": {
            "manual": len(submissions),
            "ctgov_imported": ctgov_imported,
            "ctgov_enriched": ctgov_enriched,
            "ctgov_curated": ctgov_curated,
            "pubmed_imported": pubmed_imported,
            "pubmed_enriched": pubmed_enriched,
            "pubmed_curated": pubmed_curated,
            "graph_total": len(submissions) + ctgov_curated + pubmed_curated,
            "library_total": len(submissions) + len(trials) + len(articles)
        },
        "label_map": {
            "manual": SOURCE_LABELS["manual"],
            "ctgov_imported": SOURCE_LABELS["ctgov_imported"],
            "ctgov_enriched": SOURCE_LABELS["ctgov_enriched"],
            "ctgov_curated": SOURCE_LABELS["ctgov_curated"],
            "pubmed_imported": SOURCE_LABELS["pubmed_imported"],
            "pubmed_enriched": SOURCE_LABELS["pubmed_enriched"],
            "pubmed_curated": SOURCE_LABELS["pubmed_curated"]
        }
    })


if __name__ == "__main__":
    app.run(debug=True)
